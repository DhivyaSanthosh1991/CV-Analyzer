import os, json, hmac, hashlib, uuid, re, io, base64
from flask import Flask, request, jsonify, send_file, render_template
from flask_cors import CORS
from flask_jwt_extended import JWTManager, jwt_required, get_jwt_identity, verify_jwt_in_request
import anthropic
import razorpay
from pdf_generator import generate_report_pdf
from email_sender import send_report_email
from database import (init_db, migrate_db, init_sessions_table,
                       session_set, session_get,
                       save_report, get_user_reports, get_report_pdf,
                       save_roadmap_progress, toggle_roadmap_item, get_roadmap_progress,
                       save_cv_upload, save_cv_text, get_user_cv_uploads,
                       get_cv_upload, delete_cv_upload, get_cv_storage_stats,
                       link_items_to_user, _conn as get_db_ctx)
from auth import auth_bp
from admin import admin_bp

app = Flask(__name__)
CORS(app)

# ── Config ────────────────────────────────────────────────────────────────────
ANTHROPIC_API_KEY   = os.getenv("ANTHROPIC_API_KEY",  "YOUR_ANTHROPIC_KEY")
RAZORPAY_KEY_ID     = os.getenv("RAZORPAY_KEY_ID",    "YOUR_RAZORPAY_KEY_ID")
RAZORPAY_KEY_SECRET = os.getenv("RAZORPAY_KEY_SECRET", "YOUR_RAZORPAY_SECRET")
SENDGRID_API_KEY    = os.getenv("SENDGRID_API_KEY",    "YOUR_SENDGRID_API_KEY")
JWT_SECRET          = os.getenv("JWT_SECRET_KEY",      "workmoat-dev-secret-change-in-prod")
ADMIN_SECRET        = os.getenv("ADMIN_SECRET_KEY",   "workmoat-admin-change-this")
REPORT_PRICE_PAISE  = 19900

app.config["JWT_SECRET_KEY"]       = JWT_SECRET
app.config["JWT_ACCESS_TOKEN_EXPIRES"] = False  # no expiry — user stays logged in

jwt_manager = JWTManager(app)
app.register_blueprint(auth_bp)
app.register_blueprint(admin_bp)

anthropic_client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
razorpay_client  = razorpay.Client(auth=(RAZORPAY_KEY_ID, RAZORPAY_KEY_SECRET))


# Init DB on startup
with app.app_context():
    try:
        init_db()
        migrate_db()
        init_sessions_table()
        print("✓ Database initialised successfully")
    except Exception as e:
        print(f"✗ Database init error: {e}")
        print("App will start but DB features won't work until connection is fixed")

# ── System prompt ─────────────────────────────────────────────────────────────
SYSTEM_PROMPT = """You are WorkMoat's AI Diagnostic Engine. Analyze the CV provided and return ONLY valid JSON — no markdown, no backticks, no preamble, no HTML.

SCORING RUBRICS (score each 0-100):

CV_QUALITY — Weighted composite:
Contact completeness ×0.15 | Summary/value-prop ×0.20 | Experience quantification ×0.25 | Skills currency & relevance ×0.20 | Formatting/ATS readiness ×0.20

AI_SUSCEPTIBILITY — Score the tasks, not the job title. High = repetitive, rule-based, data-processing, pattern-matching. Low = creativity, judgment, emotional intelligence, physical presence.

AI_AUGMENT_POTENTIAL — How much could AI multiply this person's output? High = mature AI tooling in domain + skills position them to adopt it.

JOB_FIT — Score against role inferred from CV. If TARGET_ROLE is provided, score against that instead.

AUTOMATION_RISK — Realistic % of listed tasks automatable within 3-5 years. Distinct from Susceptibility: seniority and strategic positioning can lower risk even in a high-susceptibility role.

TONE: Write diagnostically. If risk is HIGH or CRITICAL, be direct and frank — name the cost of inaction. Never flatter to comfort.

Return ONLY this JSON structure:
{
  "name": "<full name from CV or 'Professional'>",
  "role": "<current/target role>",
  "overall_score": <CV quality 0-100>,
  "ai_susceptibility_score": <0-100>,
  "ai_augment_score": <0-100>,
  "job_fit_score": <0-100>,
  "automation_risk": {
    "level": "<managed|moderate|high|critical>",
    "score": <0-100>,
    "timeline": "<e.g. '2-3 years'>",
    "at_risk_tasks": ["<task 1>", "<task 2>", "<task 3>"]
  },
  "cv_breakdown": {
    "contact": <0-100>,
    "summary": <0-100>,
    "experience": <0-100>,
    "skills": <0-100>,
    "formatting": <0-100>
  },
  "strengths": ["<strength 1>", "<strength 2>", "<strength 3>"],
  "gaps": ["<gap 1>", "<gap 2>", "<gap 3>"],
  "career_moat": {
    "core_strength": "<1 sentence — irreplaceable hard-to-automate value>",
    "the_threat": "<1 sentence — most at-risk part of role and why>",
    "one_move": "<1 specific actionable step with named resource or tool>"
  },
  "ai_tools_replacing": ["<tool/task 1>", "<tool/task 2>"],
  "ai_tools_to_adopt": ["<tool 1>", "<tool 2>", "<tool 3>"],
  "upskilling_roadmap": [
    {"skill": "<skill name>", "priority": "high|medium|low", "why": "<1 sentence>", "resources": "<Platform — Course name>"},
    {"skill": "<skill name>", "priority": "high|medium|low", "why": "<1 sentence>", "resources": "<Platform — Course name>"},
    {"skill": "<skill name>", "priority": "high|medium|low", "why": "<1 sentence>", "resources": "<Platform — Course name>"},
    {"skill": "<skill name>", "priority": "high|medium|low", "why": "<1 sentence>", "resources": "<Platform — Course name>"},
    {"skill": "<skill name>", "priority": "high|medium|low", "why": "<1 sentence>", "resources": "<Platform — Course name>"}
  ],
  "cv_improvements": [
    {"priority": "high", "action": "<specific CV fix — name section and what to rewrite>"},
    {"priority": "high", "action": "<specific CV fix>"},
    {"priority": "medium", "action": "<specific CV fix>"},
    {"priority": "medium", "action": "<specific CV fix>"},
    {"priority": "low", "action": "<specific CV fix>"}
  ],
  "action_plan": {
    "days_1_30": [
      "<CV fix: most urgent specific change>",
      "<AI tool adoption: specific tool for specific daily task>",
      "<AI concept: Prompt Engineering — specific free resource>",
      "<AI awareness: identify 2 tools already active in your industry>"
    ],
    "days_31_60": [
      "<AI course start: platform, course title, first module>",
      "<AI workflow: build end-to-end AI-assisted workflow for recurring task>",
      "<AI concept: Agentic AI — map one repetitive task using Make.com/Zapier/n8n>",
      "<Role signal: specific repositioning move or project to lead>"
    ],
    "days_61_90": [
      "<AI credential: complete HIGH priority course, add to LinkedIn>",
      "<AI concept: domain-specific AI application with real work output>",
      "<AI portfolio: shareable deliverable demonstrating AI fluency>",
      "<Strategic move: next role to target or application to submit>"
    ]
  },
  "strategic_direction": "<2-3 sentences on recommended role evolution for next 12-18 months>",
  "human_edge": ["<AI-resistant capability 1>", "<AI-resistant capability 2>", "<AI-resistant capability 3>"]
}
"""

# ── Helper: get optional current user ────────────────────────────────────────
def get_current_user_id():
    try:
        verify_jwt_in_request(optional=True)
        uid = get_jwt_identity()
        return int(uid) if uid else None
    except Exception:
        return None

# ── File extraction ───────────────────────────────────────────────────────────
def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    fname = filename.lower()
    if fname.endswith(".pdf"):
        try:
            import pdfplumber
            parts = []
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t: parts.append(t)
            return "\n\n".join(parts)
        except Exception as e:
            return f"[PDF extraction error: {e}]"
    elif fname.endswith(".docx"):
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            return f"[DOCX extraction error: {e}]"
    else:
        try:    return file_bytes.decode("utf-8")
        except: return file_bytes.decode("latin-1", errors="replace")

# ── Routes ────────────────────────────────────────────────────────────────────
@app.route("/")
def index():
    return render_template("index.html", razorpay_key=RAZORPAY_KEY_ID)

@app.route("/api/extract-text", methods=["POST"])
def extract_text():
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400
    f = request.files["file"]
    if not f.filename:
        return jsonify({"error": "Empty filename"}), 400

    file_bytes = f.read()
    text = extract_text_from_file(file_bytes, f.filename)
    if not text or len(text.strip()) < 20:
        return jsonify({"error": "Could not extract text from file"}), 422

    # Save CV for ALL users — link to user later if they sign in
    session_id = request.form.get("session_id") or str(uuid.uuid4())
    job_title  = request.form.get("job_title", "")
    industry   = request.form.get("industry", "")
    user_id    = get_current_user_id()

    # Always save the file — even for anonymous users
    cv_upload_id = save_cv_upload(
        session_id=session_id,
        file_bytes=file_bytes,
        filename=f.filename,
        extracted_text=text,
        job_title=job_title,
        industry=industry,
        user_id=user_id  # None if not logged in — linked later on sign-in
    )

    return jsonify({
        "text":          text[:6000],
        "session_id":    session_id,
        "cv_upload_id":  cv_upload_id
    })

@app.route("/api/analyse", methods=["POST"])
def analyse():
    data         = request.get_json()
    cv_text      = (data.get("cv_text") or "").strip()
    job_title    = (data.get("job_title") or "").strip()
    industry     = (data.get("industry") or "").strip()
    cv_upload_id = data.get("cv_upload_id")   # passed from frontend if file was uploaded

    if len(cv_text) < 60:
        return jsonify({"error": "CV text too short"}), 400

    user_msg = f"CV:\n{cv_text[:4000]}"
    if job_title: user_msg += f"\n\nTarget role: {job_title}"
    if industry:  user_msg += f"\nIndustry: {industry}"

    try:
        message = anthropic_client.messages.create(
            model="claude-sonnet-4-20250514", max_tokens=2500,
            system=SYSTEM_PROMPT,
            messages=[{"role": "user", "content": user_msg}]
        )
        raw    = message.content[0].text
        result = json.loads(re.sub(r"```json|```", "", raw).strip())
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    session_id = str(uuid.uuid4())
    session_set(session_id, result=result, paid=False)

    user_id = get_current_user_id()

    # If CV was pasted (no file upload), save the text for everyone
    if not cv_upload_id:
        cv_upload_id = save_cv_text(
            session_id=session_id,
            extracted_text=cv_text,
            job_title=job_title,
            industry=industry,
            user_id=user_id
        )

    # If user is logged in, link any existing anonymous uploads for this session
    if user_id:
        link_items_to_user(session_id, user_id)

    # Save report linked to cv_upload and user
    save_report(session_id, result, paid=False,
                user_id=user_id, cv_upload_id=cv_upload_id)

    # Build preview — all fields the frontend needs for free report
    preview = {
        "name":                   result.get("name", "Professional"),
        "role":                   result.get("role", ""),
        "overall_score":          result.get("overall_score", 0),
        "ai_susceptibility_score":result.get("ai_susceptibility_score", 0),
        "ai_augment_score":       result.get("ai_augment_score", 0),
        "job_fit_score":          result.get("job_fit_score", 0),
        "automation_risk":        result.get("automation_risk", {}),
        "cv_breakdown":           result.get("cv_breakdown", {}),
        "strengths":              result.get("strengths", []),
        "gaps":                   result.get("gaps", []),
        "career_moat":            result.get("career_moat", {}),
        "ai_tools_replacing":     result.get("ai_tools_replacing", []),
        "ai_tools_to_adopt":      result.get("ai_tools_to_adopt", []),
        "human_edge":             result.get("human_edge", []),
        "strategic_direction":    result.get("strategic_direction", ""),
    }

    return jsonify({"session_id": session_id, "preview": preview})

@app.route("/api/create-order", methods=["POST"])
def create_order():
    data       = request.get_json()
    session_id = data.get("session_id")
    email      = (data.get("email") or "").strip().lower()
    sess = session_get(session_id)
    if not sess:
        return jsonify({"error": "Session expired — please re-analyse your CV"}), 400
    if not email or "@" not in email:
        return jsonify({"error": "Valid email required"}), 400
    session_set(session_id, email=email)
    try:
        order = razorpay_client.order.create({
            "amount": REPORT_PRICE_PAISE, "currency": "INR",
            "receipt": f"wm_{session_id[:8]}",
            "notes":   {"session_id": session_id, "email": email}
        })
        session_set(session_id, order_id=order["id"])
        return jsonify({"order_id": order["id"], "amount": REPORT_PRICE_PAISE,
                        "currency": "INR", "key": RAZORPAY_KEY_ID})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/verify-payment", methods=["POST"])
def verify_payment():
    data       = request.get_json()
    session_id = data.get("session_id")
    payment_id = data.get("razorpay_payment_id")
    order_id   = data.get("razorpay_order_id")
    signature  = data.get("razorpay_signature")
    if not all([session_id, payment_id, order_id, signature]):
        return jsonify({"error": "Missing payment data"}), 400

    body         = f"{order_id}|{payment_id}"
    expected_sig = hmac.new(RAZORPAY_KEY_SECRET.encode(), body.encode(), hashlib.sha256).hexdigest()
    if not hmac.compare_digest(expected_sig, signature):
        return jsonify({"error": "Payment verification failed"}), 400

    session = session_get(session_id)
    if not session:
        return jsonify({"error": "Session not found — please re-analyse your CV"}), 404

    session_set(session_id, paid=True)
    full_result = session["result"]

    # Generate PDF
    reports_dir = os.path.join(os.path.dirname(__file__), "reports")
    os.makedirs(reports_dir, exist_ok=True)
    pdf_path = os.path.join(reports_dir, f"{session_id}.pdf")
    generate_report_pdf(full_result, pdf_path)

    pdf_b64 = ""
    try:
        with open(pdf_path, "rb") as f:
            pdf_b64 = base64.b64encode(f.read()).decode("utf-8")
    except Exception:
        pass
    session_set(session_id, pdf_b64=pdf_b64)

    # Send email — use registered email if available
    name       = full_result.get("name","Professional")
    role       = full_result.get("role","–")
    email      = session.get("email","")
    email_sent = False

    # Try to get registered email for logged-in user
    user_id_for_email = get_current_user_id()
    if user_id_for_email:
        from database import get_user_by_id
        u = get_user_by_id(user_id_for_email)
        if u and u.get("email"):
            email = u["email"]  # Use registered email

    if email and pdf_path and os.path.exists(pdf_path):
        try:
            ok, _ = send_report_email(email, name, role, full_result, pdf_path)
            email_sent = ok
        except Exception as e:
            print(f"Email send error: {e}")
            email_sent = False

    # Save paid report & link to user
    user_id    = get_current_user_id()
    report_id  = save_report(session_id, full_result, paid=True, pdf_b64=pdf_b64, user_id=user_id)

    # Auto-create roadmap progress items
    roadmap = full_result.get("upskilling_roadmap", [])
    if user_id and report_id and roadmap:
        save_roadmap_progress(user_id, report_id, roadmap)

    paid_sections = {k: full_result[k] for k in [
        "upskilling_roadmap","cv_improvements","action_plan",
        "strategic_direction","human_edge","career_moat",
        "ai_tools_replacing","ai_tools_to_adopt"
    ] if k in full_result}

    return jsonify({
        "success":       True,
        "paid_sections": paid_sections,
        "download_url":  f"/api/download/{session_id}",
        "pdf_b64":       pdf_b64,
        "pdf_name":      f"WorkMoat_Report_{name.replace(' ','_')}.pdf",
        "email_sent":    email_sent,
        "email":         email,
        "report_id":     report_id
    })

@app.route("/api/download/<session_id>")
def download_report(session_id):
    session = session_get(session_id)
    if not session or not session.get("paid"):
        return jsonify({"error": "Unauthorized"}), 403
    if not session.get("pdf_b64"):
        return jsonify({"error": "PDF not available"}), 404
    name = (session.get("result") or {}).get("name", "report").replace(" ", "_")
    return jsonify({"pdf_b64": session["pdf_b64"],
                    "pdf_name": f"WorkMoat_Report_{name}.pdf"})

@app.route("/api/user/report/<session_id>/pdf", methods=["GET"])
@jwt_required()
def user_report_pdf(session_id):
    user_id = int(get_jwt_identity())
    row = get_report_pdf(session_id, user_id)
    if not row:
        return jsonify({"error": "Report not found or not paid"}), 404
    return jsonify({"pdf_b64": row["pdf_b64"], "name": row["name"]})

@app.route("/api/user/roadmap/<int:report_id>", methods=["GET"])
@jwt_required()
def user_roadmap(report_id):
    user_id = int(get_jwt_identity())
    items   = get_roadmap_progress(user_id, report_id)
    return jsonify(items)

@app.route("/api/user/roadmap/<int:item_id>/toggle", methods=["POST"])
@jwt_required()
def toggle_roadmap(item_id):
    user_id   = int(get_jwt_identity())
    completed = request.get_json().get("completed", False)
    toggle_roadmap_item(user_id, item_id, completed)
    return jsonify({"success": True})

@app.route("/api/user/link-report", methods=["POST"])
@jwt_required()
def link_report():
    """Link anonymous session report + CV upload to the logged-in user."""
    user_id    = int(get_jwt_identity())
    session_id = request.get_json().get("session_id")
    if not session_id: return jsonify({"error": "Missing session_id"}), 400
    link_items_to_user(session_id, user_id)
    return jsonify({"success": True})


# ── CV Storage API ────────────────────────────────────────────────────────────


@app.route("/api/user/reports", methods=["GET"])
@jwt_required()
def user_reports():
    user_id = int(get_jwt_identity())
    reports = get_user_reports(user_id)
    return jsonify(reports)

@app.route("/api/user/cvs", methods=["GET"])
@jwt_required()
def list_cvs():
    user_id = int(get_jwt_identity())
    uploads = get_user_cv_uploads(user_id)
    stats   = get_cv_storage_stats(user_id)
    return jsonify({"uploads": uploads, "stats": stats})


@app.route("/api/user/cvs/<int:upload_id>", methods=["GET"])
@jwt_required()
def get_cv(upload_id):
    user_id = int(get_jwt_identity())
    cv = get_cv_upload(upload_id, user_id)
    if not cv:
        return jsonify({"error": "CV not found"}), 404
    return jsonify(cv)


@app.route("/api/user/cvs/<int:upload_id>/download", methods=["GET"])
@jwt_required()
def download_cv(upload_id):
    user_id = int(get_jwt_identity())
    cv = get_cv_upload(upload_id, user_id)
    if not cv:
        return jsonify({"error": "CV not found"}), 404
    if not cv.get("file_b64"):
        return jsonify({"error": "No file stored for this CV"}), 404

    file_bytes = base64.b64decode(cv["file_b64"])
    ext        = cv.get("file_type", "txt")
    mime_map   = {"pdf": "application/pdf",
                  "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                  "txt": "text/plain"}
    mime = mime_map.get(ext, "application/octet-stream")
    filename = cv.get("filename") or f"cv_{upload_id}.{ext}"

    return send_file(
        io.BytesIO(file_bytes),
        as_attachment=True,
        download_name=filename,
        mimetype=mime
    )


@app.route("/api/user/cvs/<int:upload_id>", methods=["DELETE"])
@jwt_required()
def delete_cv(upload_id):
    user_id = int(get_jwt_identity())
    deleted = delete_cv_upload(upload_id, user_id)
    if not deleted:
        return jsonify({"error": "CV not found or already deleted"}), 404
    return jsonify({"success": True})



@app.route("/health")
def health():
    """Health check — shows DB connection status."""
    import os
    status = {
        "app": "ok",
        "db_type": "postgresql" if os.getenv("DATABASE_URL") else "sqlite",
        "db_url_set": bool(os.getenv("DATABASE_URL")),
    }
    try:
        from database import _q
        result = _q("SELECT COUNT(*) as c FROM users", one=True)
        status["db_connected"] = True
        status["users_count"] = (result or {}).get("c", 0)
        if os.getenv("DATABASE_URL"):
            tables = _q("""SELECT table_name FROM information_schema.tables 
                          WHERE table_schema='public'""", many=True)
            status["tables"] = [t.get("table_name") for t in (tables or [])]
        status["db_status"] = "healthy"
    except Exception as e:
        status["db_connected"] = False
        status["db_error"] = str(e)
        status["db_status"] = "error"
    return jsonify(status)

if __name__ == "__main__":
    app.run(debug=True, port=5000)
